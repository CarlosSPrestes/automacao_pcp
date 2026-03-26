import streamlit as st
import pandas as pd
import re
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extrair_info(texto, campo):
    if pd.isna(texto): return "-"
    match = re.search(rf"{campo}:\s*(.*?)(?:<br />|$)", str(texto))
    return match.group(1).strip() if match else "-"

def extrair_var_nome(texto, campo):
    if pd.isna(texto): return "-"
    match = re.search(rf"<strong>{campo}</strong>\s*([^<(]+)", str(texto))
    return match.group(1).strip() if match else "-"

st.set_page_config(page_title="Gerador de Produção", page_icon="📦")
st.title("📦 Gerador de Lista de Produção")
st.write("Faça o upload do arquivo de vendas diário (CSV) para gerar o documento Word.")

# Upload do arquivo pela operadora
arquivo_csv = st.file_uploader("Arraste o arquivo pedidos.csv aqui", type=["csv"])

if arquivo_csv:
    if st.button("🚀 Gerar Documento de Produção"):
        with st.spinner("Processando dados..."):
            try:
                # 1. Ler planilhas (A base fixa deve estar na mesma pasta do app)
                df_pedidos = pd.read_csv(arquivo_csv, sep=None, engine='python', encoding='latin1')
                df_pedidos.columns = df_pedidos.columns.str.strip()

                df_base = pd.read_excel('base_produtos.xlsx', sheet_name=0)
                df_base.columns = df_base.columns.str.strip()

                # 2. Extrações
                df_pedidos['Bordado do Nome'] = df_pedidos['Informação adicional'].apply(lambda x: extrair_info(x, 'Bordado do Nome'))
                df_pedidos['Bordado Curso/Profissão'] = df_pedidos['Informação adicional'].apply(lambda x: extrair_info(x, 'Bordado Curso/Profissão'))
                df_pedidos['Cor do Bordado'] = df_pedidos['Informação adicional'].apply(lambda x: extrair_info(x, 'Cor do Bordado'))
                df_pedidos['Bolso do Busto'] = df_pedidos['Informação adicional'].apply(lambda x: extrair_info(x, 'Bolso do Busto'))
                df_pedidos['Calça Jogger'] = df_pedidos['Informação adicional'].apply(lambda x: extrair_info(x, 'Calça Jogger'))
                
                df_pedidos['Nome Limpo'] = df_pedidos['Nome produto'].apply(lambda x: str(x).split('<br')[0].strip() if pd.notna(x) else "-")
                df_pedidos['Tamanho Blusa'] = df_pedidos['Nome produto'].apply(lambda x: extrair_var_nome(x, 'Tamanho da Blusa'))
                df_pedidos['Tamanho Calça'] = df_pedidos['Nome produto'].apply(lambda x: extrair_var_nome(x, 'Tamanho da Calça'))
                df_pedidos['Tamanho Unico'] = df_pedidos['Nome produto'].apply(lambda x: extrair_var_nome(x, 'Tamanho'))

                # 3. Cruzamento
                df_base_unica = df_base.drop_duplicates(subset=['Código produto'], keep='first')
                df_base_limpa = df_base_unica.drop(columns=['Nome produto', 'Referência'], errors='ignore')
                df_final = pd.merge(df_pedidos, df_base_limpa, on='Código produto', how='left')

                # 4. Gerar Word em Memória
                doc = Document()
                
                # Resumo
                titulo_resumo = doc.add_heading('📦 RESUMO DE SEPARAÇÃO (PICKING)', level=1)
                titulo_resumo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                resumo = df_final.groupby(['Nome Limpo', 'Tamanho Blusa', 'Tamanho Calça', 'Tamanho Unico'])['Quantidade'].sum().reset_index()
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text, hdr_cells[3].text = 'Produto', 'Tam. Blusa/Único', 'Tam. Calça', 'Qtd'
                for _, row in resumo.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row['Nome Limpo'])
                    row_cells[1].text = str(row['Tamanho Blusa'] if row['Tamanho Blusa'] != '-' else row['Tamanho Unico'])
                    row_cells[2].text = str(row['Tamanho Calça'])
                    row_cells[3].text = str(row['Quantidade'])
                doc.add_page_break()

                # Lista Detalhada
                titulo_lista = doc.add_heading('📋 LISTA DE PRODUÇÃO DETALHADA', level=1)
                titulo_lista.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pedidos = df_final.groupby('Código pedido')
                for pedido_id, dados_pedido in pedidos:
                    qtd_total = int(dados_pedido['Quantidade'].sum())
                    p_pedido = doc.add_paragraph()
                    run_pedido = p_pedido.add_run(f'\n🚨 PEDIDO: #{pedido_id} | Total de Peças: {qtd_total}')
                    run_pedido.bold, run_pedido.font.size = True, Pt(14)
                    
                    contador_item = 1
                    for _, row in dados_pedido.iterrows():
                        for _ in range(int(row['Quantidade'])):
                            p_item = doc.add_paragraph()
                            p_item.add_run(f"📦 ITEM {contador_item} de {qtd_total}\n").bold = True
                            p_item.add_run(f"Produto: ").bold = True
                            p_item.add_run(f"{row.get('Nome Limpo', '-')}\n")
                            p_item.add_run(f"Tamanhos: ").bold = True
                            p_item.add_run(f"Blusa: {row.get('Tamanho Blusa', '-')} | Calça: {row.get('Tamanho Calça', '-')} | Único: {row.get('Tamanho Unico', '-')}\n")
                            p_item.add_run(f"Tecido/Cor: ").bold = True
                            p_item.add_run(f"{row.get('Tecido', '-')} / {row.get('Cor', '-')}\n")
                            p_item.add_run(f"Bordado Nome: ").bold = True
                            run_bordado = p_item.add_run(f"{row.get('Bordado do Nome', '-')}\n")
                            run_bordado.font.color.rgb = RGBColor(211, 84, 0)
                            p_item.add_run(f"Profissão: ").bold = True
                            p_item.add_run(f"{row.get('Bordado Curso/Profissão', '-')}\n")
                            p_item.add_run(f"Cor do Bordado: ").bold = True
                            p_item.add_run(f"{row.get('Cor do Bordado', '-')}\n")
                            p_item.add_run(f"Bolso do Busto: ").bold = True
                            p_item.add_run(f"{row.get('Bolso do Busto', '-')}\n")
                            p_item.add_run(f"Calça Jogger: ").bold = True
                            p_item.add_run(f"{row.get('Calça Jogger', '-')}\n")
                            p_item.add_run("-" * 40)
                            contador_item += 1

                # Salvar em buffer para download
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.success("✅ Documento gerado com sucesso!")
                st.download_button(
                    label="📥 Baixar Lista de Produção (Word)",
                    data=buffer,
                    file_name="lista_producao.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Erro ao processar: {e}")
